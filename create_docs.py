#!/usr/bin/env python3
"""
Tzniut Website Content Documents Generator
Creates two Word documents:
1. FOUNDER version - Essential info only
2. DEVELOPER version - Complete reference
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_founder_doc():
    """Create the FOUNDER version - essential info only"""
    doc = Document()
    
    # Title
    title = doc.add_heading('TZNIUT WEBSITE CONTENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Data Collection Form - Founder Edition')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(114, 47, 55)
    
    # Instructions
    doc.add_paragraph()
    doc.add_heading('Instructions', 1)
    
    para = doc.add_paragraph()
    para.add_run('IMPORTANT: ').bold = True
    para.add_run('This form contains ONLY the essential information that you (the founder) must provide. Fields marked with [REQUIRED] must be filled. Your web developer will handle all technical details and optional content.')
    
    doc.add_paragraph()
    
    # Section 1: Brand Essentials
    doc.add_heading('SECTION 1: BRAND ESSENTIALS', 1)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Your Answer'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    brand_fields = [
        ('Brand Name [REQUIRED]', ''),
        ('Tagline/Slogan [REQUIRED]', ''),
        ('Short Brand Description (1-2 sentences) [REQUIRED]', ''),
        ('WhatsApp Business Link [REQUIRED]', 'Example: https://wa.me/1234567890'),
        ('Instagram URL [REQUIRED]', ''),
        ('Facebook URL [REQUIRED]', ''),
    ]
    
    for field, placeholder in brand_fields:
        row = table.add_row().cells
        row[0].text = field
        row[1].text = placeholder if placeholder else ''
    
    doc.add_paragraph()
    
    # Section 2: Founder Info
    doc.add_heading('SECTION 2: FOUNDER INFORMATION', 1)
    
    para = doc.add_paragraph()
    para.add_run('Photo Required: ').bold = True
    para.add_run('Please provide a professional headshot (square format, minimum 400x400px)')
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Your Answer'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    founder_fields = [
        ('Full Name [REQUIRED]', ''),
        ('Title/Role [REQUIRED]', 'Example: Co-Founder & Lead Designer'),
        ('Years of Experience [REQUIRED]', ''),
        ('Professional Bio (3-4 sentences) [REQUIRED]', ''),
    ]
    
    for field, placeholder in founder_fields:
        row = table.add_row().cells
        row[0].text = field
        row[1].text = placeholder
    
    doc.add_paragraph()
    
    # Section 3: Co-Founder
    doc.add_heading('SECTION 3: CO-FOUNDER (if applicable)', 1)
    
    para = doc.add_paragraph()
    para.add_run('Note: ').italic = True
    para.add_run('Currently shows "John Doe" as placeholder. Please either provide real information or write "REMOVE" to delete this section.')
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Your Answer'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    cofounder_fields = [
        ('Full Name or "REMOVE"', ''),
        ('Title/Role (if keeping)', ''),
        ('Bio (if keeping)', ''),
        ('Photo file (if keeping)', ''),
    ]
    
    for field, _ in cofounder_fields:
        row = table.add_row().cells
        row[0].text = field
        row[1].text = ''
    
    doc.add_paragraph()
    
    # Section 4: Products
    doc.add_heading('SECTION 4: PRODUCTS - 6 Collections [REQUIRED]', 1)
    
    para = doc.add_paragraph()
    para.add_run('CRITICAL: ').bold = True
    para.add_run('Each product needs: 2 high-quality photos + price + size/color info')
    
    doc.add_paragraph()
    
    products = [
        'Velvet Drape Gown (Evening wear)',
        'Rosewood Wrap Dress (Everyday/casual)',
        'Midnight Silk Ensemble (Formal)',
        'Ivory Eternal Gown (Bridal)',
        'Champagne Soiree Dress (Evening)',
        'Power Tailored Set (Professional)',
    ]
    
    for i, name in enumerate(products, 1):
        doc.add_heading(f'4.{i} {name}', 2)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Field'
        hdr_cells[1].text = 'Your Answer'
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        product_fields = [
            ('Price or Price Range [REQUIRED]', 'Example: $450-$650'),
            ('Available Sizes', 'Example: XS-XXL'),
            ('Available Colors', ''),
            ('Materials/Fabrics', ''),
            ('Photo 1 (front view) [REQUIRED]', ''),
            ('Photo 2 (detail/alternate) [REQUIRED]', ''),
        ]
        
        for field, placeholder in product_fields:
            row = table.add_row().cells
            row[0].text = field
            row[1].text = placeholder if placeholder else ''
        
        doc.add_paragraph()
    
    # Section 5: Testimonials
    doc.add_heading('SECTION 5: CUSTOMER TESTIMONIALS [REQUIRED]', 1)
    
    para = doc.add_paragraph('Please provide 3 real testimonials with name and city.')
    
    for i in range(1, 4):
        doc.add_heading(f'Testimonial {i}', 2)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Field'
        hdr_cells[1].text = 'Your Answer'
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        row = table.add_row().cells
        row[0].text = 'Quote [REQUIRED]'
        row[1].text = ''
        row = table.add_row().cells
        row[0].text = 'Name [REQUIRED]'
        row[1].text = ''
        row = table.add_row().cells
        row[0].text = 'Location [REQUIRED]'
        row[1].text = ''
        
        doc.add_paragraph()
    
    # Section 6: Images
    doc.add_heading('SECTION 6: IMAGES TO PROVIDE', 1)
    
    doc.add_paragraph('Please provide these image files:').bold = True
    
    images_list = [
        'Logo (PNG with transparent background)',
        'Your headshot (professional, square, min 400x400px)',
        'Co-founder photo (if keeping)',
        'Hero image for homepage (high-quality fashion photo)',
        'About page studio/workspace photo',
        '12 Product photos (2 per product, portrait)',
        'WhatsApp QR code (generate from WhatsApp Business app)',
    ]
    
    for item in images_list:
        doc.add_paragraph(f'[ ] {item}', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Questions? Contact your web developer for help.')
    
    # Save
    doc.save('Tzniut_Content_Form_FOUNDER.docx')
    print("Founder document created: Tzniut_Content_Form_FOUNDER.docx")

def create_developer_doc():
    """Create the DEVELOPER version - complete reference"""
    doc = Document()
    
    # Title
    title = doc.add_heading('TZNIUT WEBSITE CONTENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Complete Reference Guide - Developer Edition')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(114, 47, 55)
    
    # Instructions
    doc.add_paragraph()
    doc.add_heading('Instructions', 1)
    doc.add_paragraph('This document contains ALL content fields. GREEN = Founder provides | BLUE = You handle | RED = Needs decision')
    
    doc.add_paragraph()
    
    # SECTION 1: Global
    doc.add_heading('SECTION 1: GLOBAL CONTENT', 1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Current Value'
    hdr_cells[2].text = 'Status'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    global_content = [
        ('Brand Name', 'Tzniut', 'FOUNDER'),
        ('Tagline', 'Take the Spotlight', 'FOUNDER'),
        ('Meta Description', 'Tzniut is a luxury fashion brand...', 'CAN REFINE'),
        ('WhatsApp Link', 'https://wa.link/qjvggk', 'FOUNDER'),
        ('Email', '(empty)', 'OPTIONAL'),
        ('Phone', '(empty)', 'OPTIONAL'),
        ('Instagram', '#', 'NEEDS REAL'),
        ('Facebook', '#', 'NEEDS REAL'),
        ('TikTok', '#', 'NEEDS REAL'),
        ('Pinterest', '#', 'NEEDS REAL'),
    ]
    
    for field, current, status in global_content:
        row = table.add_row().cells
        row[0].text = field
        row[1].text = current
        row[2].text = status
    
    doc.add_paragraph()
    
    # SECTION 2: Homepage
    doc.add_heading('SECTION 2: HOMEPAGE', 1)
    
    doc.add_heading('2.1 Hero Section', 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Current Value'
    hdr_cells[2].text = 'Status'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    hero = [
        ('Hero Image', 'Stock photo', 'NEEDS REAL PHOTO'),
        ('Hero Tagline', 'Where Modesty Meets Magnificence', 'FOUNDER'),
        ('Hero Title', 'Take the Spotlight', 'FOUNDER'),
        ('Hero Subtitle', 'Elegant, conservative fashion...', 'FOUNDER'),
    ]
    
    for field, current, status in hero:
        row = table.add_row().cells
        row[0].text = field
        row[1].text = current[:40] + '...' if len(current) > 40 else current
        row[2].text = status
    
    doc.add_paragraph()
    
    doc.add_heading('2.2 Philosophy Section', 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Current Value'
    hdr_cells[2].text = 'Status'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    philosophy = [
        ('Section Label', 'Our Philosophy', 'CAN KEEP'),
        ('Title', 'Fashion That Honours You', 'FOUNDER'),
        ('Description', 'At Tzniut, we believe...', 'FOUNDER'),
    ]
    
    for field, current, status in philosophy:
        row = table.add_row().cells
        row[0].text = field
        row[1].text = current[:40] + '...' if len(current) > 40 else current
        row[2].text = status
    
    doc.add_paragraph()
    
    doc.add_heading('2.3 Testimonials', 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Current Value'
    hdr_cells[2].text = 'Status'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    testimonials = [
        ('Testimonial 1', 'Amara J., Lagos', 'FOUNDER'),
        ('Testimonial 2', 'Fatimah K., Dubai', 'FOUNDER'),
        ('Testimonial 3', 'Grace O., London', 'FOUNDER'),
    ]
    
    for field, current, status in testimonials:
        row = table.add_row().cells
        row[0].text = field
        row[1].text = current
        row[2].text = status
    
    doc.add_paragraph()
    
    # SECTION 3: About
    doc.add_heading('SECTION 3: ABOUT PAGE', 1)
    
    doc.add_heading('3.1 Brand Story', 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Current Value'
    hdr_cells[2].text = 'Status'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    story = [
        ('Paragraph 1', 'Tzniut was born from...', 'FOUNDER'),
        ('Paragraph 2', "Michaella's journey began...", 'FOUNDER'),
        ('Paragraph 3', 'The name "Tzniut"...', 'FOUNDER'),
        ('Side Image', 'Stock photo', 'NEEDS REAL'),
    ]
    
    for field, current, status in story:
        row = table.add_row().cells
        row[0].text = field
        row[1].text = current[:40] + '...' if len(current) > 40 else current
        row[2].text = status
    
    doc.add_paragraph()
    
    doc.add_heading('3.2 Founders Section', 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Current Value'
    hdr_cells[2].text = 'Status'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    founders = [
        ('Founder 1: Michaella', 'Name, photo, bio', 'FOUNDER'),
        ('Founder 2: John Doe', 'PLACEHOLDER', 'DECIDE: Remove/Replace'),
    ]
    
    for field, current, status in founders:
        row = table.add_row().cells
        row[0].text = field
        row[1].text = current
        row[2].text = status
    
    doc.add_paragraph()
    
    # SECTION 4: Products
    doc.add_heading('SECTION 4: PRODUCTS (6 Collections)', 1)
    
    products = [
        ('Velvet Drape Gown', 'The Modest Edit'),
        ('Rosewood Wrap Dress', 'The Modest Edit'),
        ('Midnight Silk Ensemble', 'The Glory Line'),
        ('Ivory Eternal Gown', 'Bridal Modesty'),
        ('Champagne Soiree Dress', 'Evening Grace'),
        ('Power Tailored Set', 'The Working Muse'),
    ]
    
    for name, category in products:
        doc.add_heading(name, 2)
        para = doc.add_paragraph()
        para.add_run(f'Category: {category}\n').bold = True
        para.add_run('NEEDS: 2 photos + price + size/color/materials FROM FOUNDER')
    
    doc.add_paragraph()
    
    # SECTION 5: Hidden/Decisions
    doc.add_heading('SECTION 5: DECISIONS NEEDED', 1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Current Status'
    hdr_cells[2].text = 'Decision'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    decisions = [
        ('Video Showcase', 'Hidden (commented out)', 'Keep hidden until real video'),
        ('Co-Founder Section', 'Shows "John Doe"', 'Remove or replace'),
        ('Blog Posts', 'Generic content', 'Hide or rewrite'),
        ('Social Media Links', 'All # placeholders', 'Get real URLs'),
    ]
    
    for field, current, status in decisions:
        row = table.add_row().cells
        row[0].text = field
        row[1].text = current
        row[2].text = status
    
    doc.add_paragraph()
    
    # SECTION 6: Launch Checklist
    doc.add_heading('SECTION 6: PRE-LAUNCH CHECKLIST', 1)
    
    checklist_items = [
        'Received all images from founder',
        'Received all product info (photos, prices, details)',
        'Received 3 testimonials',
        'Decision: Remove co-founder?',
        'Decision: Hide blog until real content?',
        'Test all links',
        'Test mobile responsive',
        'Optimize all images',
        'Test WhatsApp button',
        'Proofread all text',
    ]
    
    for item in checklist_items:
        doc.add_paragraph(f'[ ] {item}', style='List Bullet')
    
    # Save
    doc.save('Tzniut_Content_Complete_Developer.docx')
    print("Developer document created: Tzniut_Content_Complete_Developer.docx")

if __name__ == '__main__':
    print("Creating Tzniut Website Content Documents...")
    print()
    create_founder_doc()
    print()
    create_developer_doc()
    print()
    print("=" * 60)
    print("BOTH DOCUMENTS CREATED SUCCESSFULLY!")
    print("=" * 60)
    print()
    print("1. Tzniut_Content_Form_FOUNDER.docx")
    print("   -> Send this to the founder to fill out")
    print()
    print("2. Tzniut_Content_Complete_Developer.docx")
    print("   -> Keep this for yourself as reference")
    print("=" * 60)
